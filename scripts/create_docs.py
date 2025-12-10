#!/usr/bin/env python3
"""
Create PDF and Word documentation from README.md
Uses multiple methods to ensure compatibility
"""

import os
import sys
import subprocess
from pathlib import Path

def read_markdown(file_path):
    """Read markdown file"""
    with open(file_path, 'r', encoding='utf-8') as f:
        return f.read()

def create_docx_with_pandoc():
    """Create Word document using pandoc"""
    try:
        result = subprocess.run(
            ['pandoc', 'README.md', '-o', 'documentation.docx', '--toc', '--toc-depth=3'],
            capture_output=True,
            text=True,
            timeout=60
        )
        if result.returncode == 0:
            print("✅ Word document created: documentation.docx")
            return True
        else:
            print(f"⚠️ pandoc error: {result.stderr}")
            return False
    except FileNotFoundError:
        print("⚠️ pandoc not found. Install: brew install pandoc")
        return False
    except Exception as e:
        print(f"⚠️ Error creating Word doc: {e}")
        return False

def create_docx_with_python():
    """Create Word document using python-docx"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        md_content = read_markdown('README.md')
        doc = Document()
        
        # Title
        title = doc.add_heading('GeoAI Assistant Pro', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add content (simplified - basic conversion)
        lines = md_content.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                doc.add_paragraph()
            elif line.startswith('# '):
                doc.add_heading(line[2:], level=1)
            elif line.startswith('## '):
                doc.add_heading(line[3:], level=2)
            elif line.startswith('### '):
                doc.add_heading(line[4:], level=3)
            elif line.startswith('```'):
                continue  # Skip code blocks for now
            elif line.startswith('- ') or line.startswith('* '):
                doc.add_paragraph(line[2:], style='List Bullet')
            else:
                doc.add_paragraph(line)
        
        doc.save('documentation.docx')
        print("✅ Word document created: documentation.docx")
        return True
    except ImportError:
        print("⚠️ python-docx not installed. Install: pip install python-docx")
        return False
    except Exception as e:
        print(f"⚠️ Error creating Word doc: {e}")
        return False

def create_pdf_with_chrome():
    """Create PDF from HTML using Chrome"""
    html_file = 'documentation.html'
    if not os.path.exists(html_file):
        print("⚠️ documentation.html not found")
        return False
    
    chrome_paths = [
        "/Applications/Google Chrome.app/Contents/MacOS/Google Chrome",
        "/Applications/Chromium.app/Contents/MacOS/Chromium",
    ]
    
    for chrome in chrome_paths:
        if os.path.exists(chrome):
            try:
                html_path = os.path.abspath(html_file)
                pdf_path = os.path.abspath('documentation.pdf')
                
                subprocess.run([
                    chrome,
                    '--headless',
                    '--disable-gpu',
                    f'--print-to-pdf={pdf_path}',
                    f'file://{html_path}'
                ], check=True, timeout=30, capture_output=True)
                
                if os.path.exists('documentation.pdf'):
                    print("✅ PDF created: documentation.pdf")
                    return True
            except Exception as e:
                continue
    
    print("⚠️ Chrome not found or PDF generation failed")
    return False

def create_pdf_with_pandoc():
    """Create PDF using pandoc"""
    try:
        # Try with different PDF engines
        engines = ['xelatex', 'pdflatex', 'wkhtmltopdf']
        
        for engine in engines:
            try:
                result = subprocess.run(
                    ['pandoc', 'README.md', '-o', 'documentation.pdf', 
                     f'--pdf-engine={engine}', '--toc', '--toc-depth=3'],
                    capture_output=True,
                    text=True,
                    timeout=60
                )
                if result.returncode == 0 and os.path.exists('documentation.pdf'):
                    print(f"✅ PDF created with {engine}: documentation.pdf")
                    return True
            except:
                continue
        
        return False
    except FileNotFoundError:
        print("⚠️ pandoc not found")
        return False
    except Exception as e:
        print(f"⚠️ Error: {e}")
        return False

def main():
    print("🚀 Creating documentation files...\n")
    print("="*60)
    
    # Change to script directory
    os.chdir(os.path.dirname(os.path.abspath(__file__)))
    
    # Create Word document
    print("\n📝 Creating Word document (documentation.docx)...")
    docx_success = create_docx_with_pandoc()
    if not docx_success:
        print("   Trying alternative method...")
        docx_success = create_docx_with_python()
    
    # Create PDF
    print("\n📄 Creating PDF (documentation.pdf)...")
    pdf_success = create_pdf_with_chrome()
    if not pdf_success:
        print("   Trying pandoc...")
        pdf_success = create_pdf_with_pandoc()
    
    # Summary
    print("\n" + "="*60)
    print("📊 Summary:")
    print("="*60)
    
    if pdf_success:
        size = os.path.getsize('documentation.pdf') / 1024
        print(f"✅ PDF: documentation.pdf ({size:.1f} KB)")
    else:
        print("⚠️ PDF: Could not generate automatically")
        print("   💡 Manual method:")
        print("      1. Open documentation.html in browser")
        print("      2. Print (Cmd+P) → Save as PDF")
        print("   Or install pandoc: brew install pandoc")
    
    if docx_success:
        size = os.path.getsize('documentation.docx') / 1024
        print(f"✅ Word: documentation.docx ({size:.1f} KB)")
    else:
        print("⚠️ Word: Could not generate automatically")
        print("   💡 Install pandoc: brew install pandoc")
        print("   Or: pip install python-docx markdown")
    
    print("="*60)

if __name__ == "__main__":
    main()

